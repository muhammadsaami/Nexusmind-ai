from docx import Document

doc = Document()

doc.add_heading('HR Policy Manual', level=0)  # Title style

doc.add_heading('1. Leave Policy', level=1)
doc.add_paragraph('Employees are entitled to 18 days of paid annual leave per calendar year. Leave must be applied for at least 3 working days in advance through the HR portal. Unused leave, up to a maximum of 10 days, can be carried forward to the next calendar year.')

doc.add_heading('1.1 Maternity Leave', level=2)
doc.add_paragraph('Female employees are entitled to 26 weeks of paid maternity leave as per the Maternity Benefit Act. Leave can be availed 8 weeks before the expected delivery date.')

doc.add_heading('1.2 Sick Leave', level=2)
doc.add_paragraph('Employees are entitled to 12 days of paid sick leave per year. A medical certificate is required for sick leave exceeding 2 consecutive days.')

doc.add_heading('2. Reimbursement Policy', level=1)
doc.add_paragraph('Employees can claim reimbursement for official travel, internet bills up to Rs. 1500 per month, and client meeting expenses. All claims must be submitted within 30 days with valid receipts.')

doc.add_heading('3. Work From Home Policy', level=1)
doc.add_paragraph('Employees may work from home up to 2 days per week with prior manager approval. Full-time WFH requires HR and department head sign-off.')

doc.save("../data/sample_docs/HR Policy Manual.docx")
print("Document created successfully with proper heading styles!")