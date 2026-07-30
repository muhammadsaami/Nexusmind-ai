from docx import Document

doc = Document("../data/sample_docs/HR Policy Manual.docx")

print(f"Total paragraphs found: {len(doc.paragraphs)}\n")

for i, para in enumerate(doc.paragraphs):
    style_name = para.style.name if para.style else "None"
    text_preview = para.text[:60].replace("\n", "\\n")
    print(f"[{i}] style='{style_name}' text='{text_preview}'")